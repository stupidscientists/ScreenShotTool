"""
文档管理模块
用于处理Word文档的创建、打开、保存等操作
"""

import os
import copy
import traceback
import datetime
from docx import Document
from docx.shared import Inches
from PyQt5.QtWidgets import QFileDialog, QMessageBox
from src.utils.logger import logger

class DocumentManager:
    """
    文档管理器类
    负责Word文档的创建、打开、保存和内容管理
    """
    
    def __init__(self, parent=None):
        """
        初始化文档管理器
        
        参数:
            parent: 父对象，通常是主窗口
        """
        self.parent = parent
        self.word_doc = None
        self.word_path = None
        self.original_paragraphs = []  # 保存打开文档时的原始段落数量
        self.original_rels_count = 0   # 保存打开文档时的原始关系数量（用于图片计数）
        logger.debug("初始化文档管理器")
    
    def create_document(self):
        """
        创建新的Word文档
        
        返回:
            bool: 创建成功返回True，否则返回False
        """
        try:
            # 创建新的Word文档
            logger.debug("尝试创建新的Word文档")
            self.word_doc = Document()
            
            # 添加标题
            self.word_doc.add_heading('屏幕截图文档', 0)
            
            # 添加创建时间
            current_time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            self.word_doc.add_paragraph(f'创建时间: {current_time}')
            
            # 询问保存位置
            file_path, _ = QFileDialog.getSaveFileName(
                self.parent, '保存Word文档', '', 'Word Documents (*.docx)'
            )
            
            if not file_path:
                logger.debug("用户取消了保存文档")
                self.word_doc = None
                return False
                
            # 确保文件扩展名正确
            if not file_path.endswith('.docx'):
                file_path += '.docx'
            
            # 检查目录是否存在
            dir_path = os.path.dirname(file_path)
            if not os.path.exists(dir_path):
                logger.error(f"保存目录不存在: {dir_path}")
                QMessageBox.critical(self.parent, '错误', f'保存目录不存在: {dir_path}')
                self.word_doc = None
                return False
                
            # 检查目录是否可写
            if not os.access(dir_path, os.W_OK):
                logger.error(f"保存目录无法写入: {dir_path}")
                QMessageBox.critical(self.parent, '错误', f'保存目录无法写入: {dir_path}')
                self.word_doc = None
                return False
                
            # 如果文件已存在，检查是否可写
            if os.path.exists(file_path) and not os.access(file_path, os.W_OK):
                logger.error(f"文件无法写入: {file_path}")
                QMessageBox.critical(self.parent, '错误', f'文件无法写入: {file_path}')
                self.word_doc = None
                return False
            
            # 保存文档
            try:
                logger.debug(f"尝试保存文档: {file_path}")
                
                # 创建临时文件名
                temp_path = file_path + ".tmp"
                
                # 先保存到临时文件
                self.word_doc.save(temp_path)
                
                # 如果原文件存在，先备份
                if os.path.exists(file_path):
                    backup_path = file_path + ".bak"
                    try:
                        if os.path.exists(backup_path):
                            os.remove(backup_path)
                        os.rename(file_path, backup_path)
                        logger.debug(f"已创建备份文件: {backup_path}")
                    except Exception as e:
                        logger.warning(f"创建备份文件时出错: {str(e)}")
                
                # 将临时文件重命名为目标文件
                os.rename(temp_path, file_path)
                
                self.word_path = file_path
                
                # 记录文件的最后修改时间
                self.last_modified_time = os.path.getmtime(file_path)
                logger.debug(f"记录文件最后修改时间: {self.last_modified_time}")
                
                # 保存原始文档内容的信息
                self.save_original_content_info()
                
                logger.info(f"文档已成功创建并保存: {file_path}")
                return True
                
            except Exception as e:
                logger.error(f"保存文档时出错: {str(e)}")
                logger.error(traceback.format_exc())
                
                # 检查临时文件是否存在，如果存在则删除
                if os.path.exists(temp_path):
                    try:
                        os.remove(temp_path)
                    except:
                        pass
                
                # 显示错误消息
                QMessageBox.critical(self.parent, '错误', f'保存Word文档时出错: {str(e)}')
                self.word_doc = None
                self.word_path = None
                return False
                
        except Exception as e:
            logger.error(f"创建Word文档时出错: {str(e)}")
            logger.error(traceback.format_exc())
            QMessageBox.critical(self.parent, '错误', f'创建Word文档时出错: {str(e)}')
            self.word_doc = None
            self.word_path = None
            return False
    
    def open_document(self):
        """
        打开现有Word文档
        
        返回:
            bool: 打开成功返回True，否则返回False
        """
        # 打开现有Word文档
        file_path, _ = QFileDialog.getOpenFileName(
            self.parent, '打开Word文档', '', 'Word Documents (*.docx)'
        )
        
        if not file_path:
            logger.debug("用户取消了打开文档操作")
            return False
            
        # 检查文件是否存在
        if not os.path.exists(file_path):
            logger.error(f"文件不存在: {file_path}")
            QMessageBox.critical(self.parent, '错误', f'文件不存在: {file_path}')
            return False
            
        # 检查文件是否可读
        if not os.access(file_path, os.R_OK):
            logger.error(f"文件无法读取: {file_path}")
            QMessageBox.critical(self.parent, '错误', f'文件无法读取: {file_path}')
            return False
            
        # 检查文件大小
        try:
            file_size = os.path.getsize(file_path)
            if file_size == 0:
                logger.error(f"文件为空: {file_path}")
                QMessageBox.critical(self.parent, '错误', f'文件为空: {file_path}')
                return False
        except Exception as e:
            logger.error(f"检查文件大小时出错: {str(e)}")
            QMessageBox.critical(self.parent, '错误', f'检查文件时出错: {str(e)}')
            return False
            
        try:
            logger.debug(f"尝试打开文档: {file_path}")
            self.word_doc = Document(file_path)
            self.word_path = file_path
            
            # 记录文件的最后修改时间
            self.last_modified_time = os.path.getmtime(file_path)
            logger.debug(f"记录文件最后修改时间: {self.last_modified_time}")
            
            # 验证文档是否有效
            try:
                _ = self.word_doc.paragraphs
                logger.debug("文档有效，可以访问段落")
                
                # 保存原始文档内容的信息
                self.save_original_content_info()
                
            except Exception as e:
                logger.error(f"文档无效: {str(e)}")
                QMessageBox.critical(self.parent, '错误', f'文档格式无效: {str(e)}')
                self.word_doc = None
                self.word_path = None
                return False
                
            logger.info(f"成功打开文档: {file_path}")
            return True
                
        except Exception as e:
            logger.error(f"打开文档时出错: {str(e)}")
            logger.error(traceback.format_exc())
            QMessageBox.critical(self.parent, '错误', f'打开文档时出错: {str(e)}')
            self.word_doc = None
            self.word_path = None
            return False
    
    def save_document(self):
        """
        保存当前Word文档
        
        返回:
            bool: 保存成功返回True，否则返回False
        """
        if not self.word_doc or not self.word_path:
            logger.warning("尝试保存文档但未创建或打开文档")
            QMessageBox.warning(self.parent, '警告', '请先创建或打开Word文档')
            return False
        
        # 检查文件路径是否有效
        try:
            # 检查目录是否存在
            dir_path = os.path.dirname(self.word_path)
            if not os.path.exists(dir_path):
                logger.error(f"保存目录不存在: {dir_path}")
                QMessageBox.critical(self.parent, '错误', f'保存目录不存在: {dir_path}')
                return False
                
            # 检查目录是否可写
            if not os.access(dir_path, os.W_OK):
                logger.error(f"保存目录无法写入: {dir_path}")
                QMessageBox.critical(self.parent, '错误', f'保存目录无法写入: {dir_path}')
                return False
                
            # 如果文件已存在，检查是否可写
            if os.path.exists(self.word_path) and not os.access(self.word_path, os.W_OK):
                logger.error(f"文件无法写入: {self.word_path}")
                QMessageBox.critical(self.parent, '错误', f'文件无法写入: {self.word_path}')
                return False
                
            # 检查文件是否被外部修改（如果文件存在）
            if os.path.exists(self.word_path) and hasattr(self, 'last_modified_time'):
                current_mtime = os.path.getmtime(self.word_path)
                
                if current_mtime != self.last_modified_time:
                    logger.warning(f"检测到文件已被外部修改: {self.word_path}")
                    
                    # 询问用户如何处理
                    msg_box = QMessageBox(self.parent)
                    msg_box.setWindowTitle('文件已修改')
                    msg_box.setText('检测到文档已在外部被修改。如何处理?')
                    msg_box.setInformativeText('【合并】: 尝试将当前内容合并到修改后的文档\n【覆盖】: 使用当前内容覆盖外部修改\n【取消】: 取消保存操作')
                    
                    # 添加自定义按钮
                    merge_button = msg_box.addButton('合并', QMessageBox.YesRole)
                    overwrite_button = msg_box.addButton('覆盖', QMessageBox.AcceptRole)
                    cancel_button = msg_box.addButton('取消', QMessageBox.RejectRole)
                    
                    # 设置默认按钮
                    msg_box.setDefaultButton(merge_button)
                    
                    # 执行对话框
                    msg_box.exec_()
                    
                    # 获取用户点击的按钮
                    clicked_button = msg_box.clickedButton()
                    
                    if clicked_button == cancel_button:
                        logger.info("用户取消了保存操作")
                        return False
                    
                    elif clicked_button == merge_button:  # 合并
                        logger.info("用户选择合并文档")
                        
                        try:
                            # 尝试合并文档
                            # 1. 先读取当前文档内容
                            current_doc = self.word_doc
                            
                            # 2. 重新加载磁盘上的文档
                            self.word_doc = Document(self.word_path)
                            
                            # 3. 确定新增内容，只合并自上次打开后添加的内容
                            # 添加一个分隔行，表明下面是新增内容
                            # separator = self.word_doc.add_paragraph("=" * 30)
                            # separator.add_run(" 以下是新增内容 ").bold = True
                            # separator.add_run("=" * 30)
                            
                            # 复制新增的段落文本（跳过原始内容）
                            original_count = len(self.original_paragraphs)
                            current_count = len(list(current_doc.paragraphs))
                            
                            logger.debug(f"原始段落数: {original_count}, 当前段落数: {current_count}")
                            
                            if current_count > original_count:
                                # 只复制新增的段落
                                logger.info(f"将复制 {current_count - original_count} 个新增段落")
                                added_paragraphs = list(current_doc.paragraphs)[original_count:]
                                
                                for paragraph in added_paragraphs:
                                    if paragraph.text:
                                        p = self.word_doc.add_paragraph()
                                        for run in paragraph.runs:
                                            # 复制文本及其格式
                                            new_run = p.add_run(run.text)
                                            new_run.bold = run.bold
                                            new_run.italic = run.italic
                                            new_run.underline = run.underline
                                            # 可以添加更多格式复制
                            else:
                                logger.info("没有检测到新增段落，将复制最后一个段落作为新增内容")
                                # 如果没有新增段落，复制最后一个段落作为提示
                                if current_count > 0:
                                    last_para = current_doc.paragraphs[-1]
                                    p = self.word_doc.add_paragraph()
                                    for run in last_para.runs:
                                        new_run = p.add_run(run.text)
                                        new_run.bold = run.bold
                                        new_run.italic = run.italic
                                        new_run.underline = run.underline
                            
                            # 使用更健壮的方式复制新增图片
                            try:
                                # 创建临时目录用于存储图片
                                app_dir = os.path.abspath(os.path.dirname(os.path.dirname(os.path.dirname(__file__))))
                                temp_dir = os.path.join(app_dir, 'temp_merge_images')
                                os.makedirs(temp_dir, exist_ok=True)
                                
                                # 复制新增图片（跳过原始图片）
                                image_count = 0
                                temp_files = []  # 记录创建的临时文件
                                
                                # 计算当前文档中的图片关系数
                                current_rels_count = len(current_doc.part.rels)
                                logger.debug(f"原始关系数: {self.original_rels_count}, 当前关系数: {current_rels_count}")
                                
                                # 只复制新增的图片关系
                                if current_rels_count > self.original_rels_count:
                                    # 尝试找出新增的图片关系
                                    all_rels = list(current_doc.part.rels.values())
                                    new_rels = all_rels[self.original_rels_count:]
                                    
                                    logger.info(f"检测到 {len(new_rels)} 个新增关系")
                                    
                                    for rel in new_rels:
                                        try:
                                            # 检查是否是图片关系
                                            if hasattr(rel, 'target_ref') and "image" in rel.target_ref.lower():
                                                image_count += 1
                                                # 获取图片二进制数据
                                                image_part = rel._target
                                                if hasattr(image_part, 'blob'):
                                                    # 保存图片到临时文件
                                                    img_ext = os.path.splitext(rel.target_ref)[1]
                                                    if not img_ext:
                                                        img_ext = ".png"  # 默认扩展名
                                                    
                                                    temp_img_path = os.path.join(
                                                        temp_dir, f'merge_image_{image_count}{img_ext}'
                                                    )
                                                    
                                                    with open(temp_img_path, 'wb') as f:
                                                        f.write(image_part.blob)
                                                    
                                                    # 将图片添加到文档
                                                    self.word_doc.add_picture(temp_img_path, width=Inches(6))
                                                    
                                                    # 添加到临时文件列表
                                                    temp_files.append(temp_img_path)
                                                    
                                                    logger.debug(f"成功复制图片 {image_count} 到合并文档")
                                        except Exception as img_error:
                                            logger.warning(f"处理图片时出错: {str(img_error)}")
                                            continue
                                else:
                                    logger.info("没有检测到新增图片关系")
                                
                                if image_count > 0:
                                    logger.info(f"成功复制 {image_count} 张新增图片到合并文档")
                                else:
                                    logger.info("未找到新增图片需要复制")
                                    
                                # 清理临时文件
                                try:
                                    for temp_file in temp_files:
                                        if os.path.exists(temp_file):
                                            try:
                                                os.remove(temp_file)
                                                logger.debug(f"已删除临时文件: {temp_file}")
                                            except Exception:
                                                logger.warning(f"无法删除临时文件: {temp_file}")
                                    
                                    # 尝试删除临时目录（如果为空）
                                    try:
                                        os.rmdir(temp_dir)
                                        logger.debug(f"已删除临时目录: {temp_dir}")
                                    except Exception:
                                        logger.debug(f"无法删除临时目录，可能不为空: {temp_dir}")
                                except Exception as cleanup_error:
                                    logger.warning(f"清理临时文件时出错: {str(cleanup_error)}")
                                    # 继续执行，不因清理失败而中止
                                    
                            except Exception as img_copy_error:
                                logger.error(f"复制图片过程中出错: {str(img_copy_error)}")
                                logger.error(traceback.format_exc())
                                # 继续执行，不因图片复制失败而中止整个合并过程
                            
                            logger.info("文档合并成功")
                            
                        except Exception as e:
                            logger.error(f"合并文档时出错: {str(e)}")
                            logger.error(traceback.format_exc())
                            
                            # 提示用户合并失败
                            merge_msg_box = QMessageBox(self.parent)
                            merge_msg_box.setWindowTitle('合并失败')
                            merge_msg_box.setText(f'合并文档失败: {str(e)}')
                            merge_msg_box.setInformativeText('是否仍要覆盖外部修改？')
                            
                            # 添加自定义按钮
                            overwrite_button = merge_msg_box.addButton('覆盖', QMessageBox.YesRole)
                            cancel_button = merge_msg_box.addButton('取消', QMessageBox.RejectRole)
                            
                            # 设置默认按钮
                            merge_msg_box.setDefaultButton(cancel_button)
                            
                            # 执行对话框
                            merge_msg_box.exec_()
                            
                            # 获取用户点击的按钮
                            if merge_msg_box.clickedButton() == cancel_button:
                                logger.info("用户取消了保存操作")
                                return False
                    
                    # 如果用户选择"覆盖"或合并失败后选择继续，则使用当前内容覆盖
                    if clicked_button == overwrite_button:
                        logger.info("用户选择覆盖外部修改")
                    else:
                        logger.info("将使用当前内容覆盖文件")
        except Exception as e:
            logger.error(f"检查文件路径时出错: {str(e)}")
            QMessageBox.critical(self.parent, '错误', f'检查文件路径时出错: {str(e)}')
            return False
        
        # 检查文档对象是否有效
        try:
            _ = self.word_doc.paragraphs
        except Exception as e:
            logger.error(f"文档对象无效: {str(e)}")
            QMessageBox.critical(self.parent, '错误', f'文档对象无效，无法保存: {str(e)}')
            return False
        
        # 尝试保存文档
        try:
            logger.debug(f"尝试保存文档: {self.word_path}")
            
            # 创建临时文件名
            temp_path = self.word_path + ".tmp"
            
            # 先保存到临时文件
            self.word_doc.save(temp_path)
            
            # 如果原文件存在，先备份
            if os.path.exists(self.word_path):
                backup_path = self.word_path + ".bak"
                try:
                    if os.path.exists(backup_path):
                        os.remove(backup_path)
                    os.rename(self.word_path, backup_path)
                    logger.debug(f"已创建备份文件: {backup_path}")
                except Exception as e:
                    logger.warning(f"创建备份文件时出错: {str(e)}")
            
            # 将临时文件重命名为目标文件
            os.rename(temp_path, self.word_path)
            
            # 更新最后修改时间记录
            self.last_modified_time = os.path.getmtime(self.word_path)
            
            # 更新原始内容快照
            self.save_original_content_info()
            
            logger.info(f"文档已成功保存: {self.word_path}")
            return True
            
        except Exception as e:
            logger.error(f"保存文档时出错: {str(e)}")
            logger.error(traceback.format_exc())
            
            # 检查临时文件是否存在，如果存在则删除
            if os.path.exists(temp_path):
                try:
                    os.remove(temp_path)
                except:
                    pass
            
            # 显示错误消息
            QMessageBox.critical(self.parent, '错误', f'保存Word文档时出错: {str(e)}')
            
            # 尝试恢复备份
            backup_path = self.word_path + ".bak"
            if os.path.exists(backup_path):
                try:
                    if os.path.exists(self.word_path):
                        os.remove(self.word_path)
                    os.rename(backup_path, self.word_path)
                    logger.info(f"已从备份恢复文档: {self.word_path}")
                    QMessageBox.information(self.parent, '恢复', '已从备份恢复文档')
                except Exception as restore_error:
                    logger.error(f"恢复备份时出错: {str(restore_error)}")
                    QMessageBox.warning(self.parent, '警告', f'无法恢复备份: {str(restore_error)}\n备份文件位于: {backup_path}')
            
            return False
    
    def add_screenshot(self, pixmap, text=""):
        """
        将截图添加到Word文档
        
        参数:
            pixmap: QPixmap对象，要添加的截图
            text: 字符串，截图的说明文本
            
        返回:
            bool: 添加成功返回True，否则返回False
        """
        logger.info("开始添加截图到Word文档")
        try:
            # 检查文档是否有效
            if not self.word_doc:
                logger.error("Word文档未打开")
                raise Exception("Word文档未打开")
            
            # 检查文档是否有效
            try:
                _ = self.word_doc.paragraphs
            except Exception as e:
                logger.error(f"Word文档无效: {str(e)}")
                
                # 尝试重新打开文档
                try:
                    logger.debug("尝试重新打开文档")
                    self.word_doc = Document(self.word_path)
                    
                    # 再次检查文档是否有效
                    _ = self.word_doc.paragraphs
                    logger.debug("重新打开文档成功")
                    
                    # 更新最后修改时间记录
                    self.last_modified_time = os.path.getmtime(self.word_path)
                    logger.debug(f"更新文件最后修改时间: {self.last_modified_time}")
                except Exception as reopen_error:
                    logger.error(f"重新打开文档失败: {str(reopen_error)}")
                    raise Exception(f"Word文档无效且无法重新打开: {str(e)}")
            
            # 检查pixmap是否有效
            if pixmap.isNull():
                logger.error("截图无效，无法保存")
                raise Exception("截图无效，无法保存")
            
            # 创建临时目录（在应用程序目录下）
            try:
                # 获取应用程序目录
                app_dir = os.path.abspath(os.path.dirname(os.path.dirname(os.path.dirname(__file__))))
                temp_dir = os.path.join(app_dir, 'temp_screenshots')
                os.makedirs(temp_dir, exist_ok=True)
                
                # 添加更详细的日志
                logger.info("=" * 50)
                logger.info(f"临时截图文件夹路径: {temp_dir}")
                logger.info(f"应用程序目录: {app_dir}")
                logger.info(f"当前文件位置: {__file__}")
                logger.info("=" * 50)
                
                logger.debug(f"创建临时目录: {temp_dir}")
            except Exception as e:
                logger.error(f"创建临时目录时出错: {str(e)}")
                raise Exception(f"创建临时目录时出错: {str(e)}")
            
            # 保存临时图片文件
            try:
                # 使用时间戳作为文件名，避免冲突
                timestamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S%f")
                temp_img_path = os.path.join(temp_dir, f'screenshot_{timestamp}.png')
                logger.debug(f"保存临时图片: {temp_img_path}")
                
                saved = pixmap.save(temp_img_path)
                if not saved:
                    logger.error(f"保存临时图片失败: {temp_img_path}")
                    raise Exception(f"保存临时图片失败: {temp_img_path}")
            except Exception as e:
                logger.error(f"保存临时图片时出错: {str(e)}")
                raise Exception(f"保存临时图片时出错: {str(e)}")
            
            # 添加截图到Word文档
            try:
                # 添加文本说明（如果有）
                if text:
                    self.word_doc.add_paragraph(text)
                
                # 添加图片
                self.word_doc.add_picture(temp_img_path, width=Inches(6))
                
                # 添加空行
                self.word_doc.add_paragraph()
                
                logger.info("成功添加截图到Word文档")
                
                # 自动保存文档
                if self.save_document():
                    # 更新最后修改时间记录
                    self.last_modified_time = os.path.getmtime(self.word_path)
                    logger.debug(f"更新文件最后修改时间: {self.last_modified_time}")
                
                return True
            except Exception as e:
                logger.error(f"添加截图到Word文档时出错: {str(e)}")
                raise Exception(f"添加截图到Word文档时出错: {str(e)}")
            
        except Exception as e:
            logger.error(f"添加截图到Word文档失败: {str(e)}")
            if self.parent:
                QMessageBox.critical(self.parent, '错误', f'添加截图到Word文档失败: {str(e)}')
            return False
    
    def close_document(self, ask_save=True):
        """
        关闭当前文档
        
        参数:
            ask_save: 是否询问用户是否保存
            
        返回:
            bool: 关闭成功返回True，否则返回False
        """
        if not self.word_doc or not self.word_path:
            logger.debug("没有打开的文档，无需关闭")
            return True
            
        if ask_save:
            # 询问用户是否保存
            reply = QMessageBox.question(
                self.parent, '保存文档', 
                '是否在关闭前保存文档？', 
                QMessageBox.Yes | QMessageBox.No | QMessageBox.Cancel
            )
            
            if reply == QMessageBox.Cancel:
                logger.info("用户取消了关闭操作")
                return False
                
            if reply == QMessageBox.Yes:
                if not self.save_document():
                    # 保存失败，询问是否继续关闭
                    reply = QMessageBox.question(
                        self.parent, '保存失败', 
                        '保存文档失败，是否仍然关闭文档？', 
                        QMessageBox.Yes | QMessageBox.No
                    )
                    
                    if reply == QMessageBox.No:
                        logger.info("用户选择取消关闭")
                        return False
        
        # 关闭文档
        self.word_doc = None
        self.word_path = None
        logger.info("文档已关闭")
        return True
    
    def save_original_content_info(self):
        """
        保存原始文档内容的信息，用于后续比较和合并
        """
        try:
            if self.word_doc:
                # 保存段落内容
                self.original_paragraphs = []
                for para in self.word_doc.paragraphs:
                    self.original_paragraphs.append(para.text)
                
                # 保存关系数量（用于图片计数）
                self.original_rels_count = len(self.word_doc.part.rels)
                
                logger.debug(f"保存原始文档信息：{len(self.original_paragraphs)}段落，{self.original_rels_count}个关系")
        except Exception as e:
            logger.error(f"保存原始文档内容信息时出错: {str(e)}")
            logger.error(traceback.format_exc()) 